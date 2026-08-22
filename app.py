"""Personal English Dictionary built with Streamlit and Google Firestore."""

from __future__ import annotations

import json
import os
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

import firebase_admin
import streamlit as st
from firebase_admin import credentials, firestore
from google.cloud.firestore_v1 import Client


st.set_page_config(page_title="My English Dictionary", page_icon="📚", layout="wide")

DEFAULT_TAGS = ["関係代名詞", "仮定法", "分詞構文", "比較", "SVO+C", "時制", "助動詞", "イディオム"]
DEFAULT_LESSON_CATEGORIES = ["導入", "解説", "問題演習", "グループワーク", "発表", "振り返り", "その他"]
READING_LEVELS = ["中学1年", "中学2年", "中学3年", "高校入試"]
READING_QUESTION_TYPES = ["語順並べ替え問題", "和文英訳問題", "内容一致問題", "空所補充問題", "下線部和訳問題", "指示語・文脈把握問題", "要約・主旨把握問題"]


def secret_or_env(name: str) -> str | None:
    """Read a non-secret setting from Streamlit Secrets first, then the environment."""
    try:
        value = st.secrets.get(name)
    except Exception:
        value = None
    return value or os.getenv(name)


@st.cache_resource(show_spinner=False)
def get_firestore_client(service_account_json: str, database_id: str | None) -> Client:
    """Initialize the Firebase Admin SDK only once, then return Firestore."""
    service_account = json.loads(service_account_json)
    if not firebase_admin._apps:
        firebase_admin.initialize_app(credentials.Certificate(service_account))
    return firestore.client(database_id=database_id)


def get_client() -> Client | None:
    """Read credentials from Secrets, an environment variable, or a local JSON file."""
    service_account: dict[str, Any] | None = None
    try:
        if "firebase_service_account" in st.secrets:
            service_account = dict(st.secrets["firebase_service_account"])
    except Exception:
        pass

    # Useful for local development. Never commit this JSON file.
    credential_path = os.getenv("FIREBASE_SERVICE_ACCOUNT_PATH")
    if service_account is None and not credential_path:
        # Firebase names downloaded Admin SDK keys like *-firebase-adminsdk-*.json.
        # The file remains local and is ignored by Git; deployment should use Secrets.
        candidates = [Path.cwd() / "firebase-service-account.json"]
        candidates.extend(Path.cwd().glob("*-firebase-adminsdk-*.json"))
        credential_path = next((str(path) for path in candidates if path.is_file()), None)

    if service_account is None and credential_path:
        try:
            with open(credential_path, encoding="utf-8") as file:
                service_account = json.load(file)
        except (OSError, json.JSONDecodeError) as exc:
            st.error(f"サービスアカウントJSONを読み込めませんでした: {exc}")
            return None

    if service_account is None:
        st.error("Firebase認証情報が未設定です。`.streamlit/secrets.toml` またはサービスアカウントJSONを設定してください。")
        return None
    try:
        database_id = secret_or_env("FIRESTORE_DATABASE_ID")
        return get_firestore_client(json.dumps(service_account), database_id)
    except Exception as exc:
        st.error(f"Firestoreへ接続できませんでした: {exc}")
        return None


def contains_keyword(row: dict[str, Any], fields: list[str], keyword: str) -> bool:
    if not keyword:
        return True
    needle = keyword.lower()
    return any(needle in str(row.get(field) or "").lower() for field in fields)


def format_date(value: Any) -> str:
    if isinstance(value, datetime):
        return value.astimezone().strftime("%Y-%m-%d %H:%M")
    return str(value).replace("T", " ")[:16] if value else ""


def show_firestore_error(action: str, exc: Exception) -> None:
    """Show an actionable message for a disabled Firestore API."""
    message = str(exc)
    if "SERVICE_DISABLED" in message or "firestore.googleapis.com" in message:
        st.error(f"{action}できませんでした: Cloud Firestore APIが有効化されていません。")
        st.markdown(
            "[FirebaseプロジェクトのCloud Firestore APIを有効化する](https://console.developers.google.com/apis/api/firestore.googleapis.com/overview?project=my-dictionary-date-base)"
        )
        st.info("有効化後、反映に数分かかることがあります。数分待ってからもう一度保存してください。")
    elif "database (default) does not exist" in message:
        st.error(f"{action}できませんでした: Firestoreデータベースがまだ作成されていません。")
        st.markdown(
            "[Firebase ConsoleでFirestoreデータベースを作成する](https://console.cloud.google.com/datastore/setup?project=my-dictionary-date-base)"
        )
        st.info("作成時は「Cloud Firestore」「Native mode」を選んでください。作成後にもう一度保存できます。")
    else:
        st.error(f"{action}できませんでした: {exc}")


def fetch_rows(client: Client, collection: str) -> list[dict[str, Any]]:
    """Fetch one collection, newest first. Filtering happens in the app for flexible search."""
    documents = client.collection(collection).order_by("created_at", direction=firestore.Query.DESCENDING).stream()
    return [{"id": document.id, **document.to_dict()} for document in documents]


def delete_document(client: Client, collection: str, document_id: str, label: str) -> None:
    """Delete only after the user confirms from the card popover."""
    try:
        client.collection(collection).document(document_id).delete()
        st.success(f"{label}を削除しました。")
        st.rerun()
    except Exception as exc:
        show_firestore_error(f"{label}を削除", exc)


def quote_list(client: Client) -> None:
    st.subheader("登録済みの引用")
    keyword = st.text_input("キーワード", placeholder="引用文・出典・メモを検索", key="quote_keyword")
    try:
        rows = fetch_rows(client, "quotes")
    except Exception as exc:
        show_firestore_error("引用データを取得", exc)
        return

    all_tags = sorted({tag for row in rows for tag in (row.get("tags") or [])})
    selected_tags = st.multiselect("文法タグで絞り込み", all_tags, key="quote_tags")
    filtered = [
        row for row in rows
        if contains_keyword(row, ["sentence", "source", "memo"], keyword)
        and set(selected_tags).issubset(set(row.get("tags") or []))
    ]
    st.caption(f"{len(filtered)} 件表示")
    if not filtered:
        st.info("該当する引用はありません。新規登録タブから追加できます。")
        return
    for row in filtered:
        with st.container(border=True):
            st.markdown(f"> {row.get('sentence', '')}")
            if row.get("translation"):
                st.write(f"**訳**　{row['translation']}")
            if row.get("source"):
                st.caption(f"出典: {row['source']}")
            if row.get("tags"):
                st.caption("　".join(f"`{tag}`" for tag in row["tags"]))
            if row.get("memo"):
                st.write(f"**メモ**　{row['memo']}")
            if row.get("created_at"):
                st.caption(f"登録: {format_date(row['created_at'])}")
            with st.popover("削除", icon="🗑️"):
                st.warning("この引用を削除します。元に戻せません。")
                if st.button("この引用を削除する", type="primary", key=f"delete_quote_{row['id']}"):
                    delete_document(client, "quotes", row["id"], "引用")


def quote_form(client: Client) -> None:
    st.subheader("引用を新規登録")
    with st.form("quote_form", clear_on_submit=True):
        sentence = st.text_area("引用文 *", placeholder="To be, or not to be...")
        translation = st.text_area("日本語訳 *")
        source = st.text_input("出典", placeholder="Hamlet / William Shakespeare")
        tags = st.multiselect("文法タグ", DEFAULT_TAGS, accept_new_options=True)
        memo = st.text_area("自分用メモ")
        submitted = st.form_submit_button("保存する", type="primary")
    if not submitted:
        return
    if not sentence.strip() or not translation.strip():
        st.warning("「引用文」と「日本語訳」は必須です。")
        return
    payload = {
        "sentence": sentence.strip(), "translation": translation.strip(), "source": source.strip(),
        "tags": tags, "memo": memo.strip(), "created_at": datetime.now(timezone.utc),
    }
    try:
        client.collection("quotes").add(payload)
        st.success("引用をFirestoreへ保存しました。")
    except Exception as exc:
        show_firestore_error("引用を保存", exc)


def etymology_list(client: Client) -> None:
    st.subheader("登録済みの語源")
    keyword = st.text_input("キーワード", placeholder="単語・語源・メモを検索", key="etymology_keyword")
    try:
        rows = fetch_rows(client, "etymologies")
    except Exception as exc:
        show_firestore_error("語源データを取得", exc)
        return
    filtered = [row for row in rows if contains_keyword(row, ["words", "root", "memo"], keyword)]
    st.caption(f"{len(filtered)} 件表示")
    if not filtered:
        st.info("該当する語源はありません。新規登録タブから追加できます。")
        return
    for row in filtered:
        with st.container(border=True):
            st.markdown(f"### {row.get('words', '')}")
            st.write(f"**語源**　{row.get('root', '')}")
            if row.get("memo"):
                st.write(row["memo"])
            if row.get("created_at"):
                st.caption(f"登録: {format_date(row['created_at'])}")
            with st.popover("削除", icon="🗑️"):
                st.warning("この語源を削除します。元に戻せません。")
                if st.button("この語源を削除する", type="primary", key=f"delete_etymology_{row['id']}"):
                    delete_document(client, "etymologies", row["id"], "語源")


def etymology_form(client: Client) -> None:
    st.subheader("語源を新規登録")
    with st.form("etymology_form", clear_on_submit=True):
        words = st.text_input("現在の単語（代表例） *", placeholder="dictionary, dictate")
        root = st.text_input("語源となる言葉 *", placeholder="dict = 言う")
        memo = st.text_area("由来・メモ", placeholder="覚え方や関連語を書き留める")
        submitted = st.form_submit_button("保存する", type="primary")
    if not submitted:
        return
    if not words.strip() or not root.strip():
        st.warning("「現在の単語」と「語源となる言葉」は必須です。")
        return
    payload = {"words": words.strip(), "root": root.strip(), "memo": memo.strip(), "created_at": datetime.now(timezone.utc)}
    try:
        client.collection("etymologies").add(payload)
        st.success("語源をFirestoreへ保存しました。")
    except Exception as exc:
        show_firestore_error("語源を保存", exc)


def fetch_lesson_cards(client: Client, lesson_id: str) -> list[dict[str, Any]]:
    """Load cards for one lesson and sort them by elapsed time in the app."""
    rows = fetch_rows(client, "lesson_cards")
    return sorted((row for row in rows if row.get("lesson_id") == lesson_id), key=lambda row: (int(row.get("elapsed_minutes") or 0), str(row.get("created_at") or "")))


def lesson_label(lesson: dict[str, Any]) -> str:
    date, title, unit = lesson.get("lesson_date") or "日付未設定", lesson.get("title") or "無題の授業", lesson.get("unit")
    return f"{date}｜{title}" + (f"（{unit}）" if unit else "")


def lesson_settings(client: Client) -> None:
    st.subheader("授業・単元の設定")
    st.caption("授業を作成すると、その授業ごとにタイムラインを記録できます。")
    with st.form("lesson_form", clear_on_submit=True):
        title = st.text_input("授業名 *", placeholder="例：英語コミュニケーションI")
        unit = st.text_input("単元名", placeholder="例：Lesson 3 Food and Culture")
        lesson_date = st.date_input("授業日", value=datetime.now().date())
        memo = st.text_area("授業メモ", placeholder="ねらい・準備物など")
        submitted = st.form_submit_button("授業を作成する", type="primary")
    if not submitted:
        return
    if not title.strip():
        st.warning("「授業名」は必須です。")
        return
    now = datetime.now(timezone.utc)
    try:
        client.collection("lessons").add({"title": title.strip(), "unit": unit.strip(), "lesson_date": lesson_date.isoformat(), "memo": memo.strip(), "created_at": now, "updated_at": now})
        st.success("授業を作成しました。タイムラインタブからカードを追加できます。")
    except Exception as exc:
        show_firestore_error("授業を作成", exc)


def timeline_card(row: dict[str, Any]) -> None:
    st.markdown(f"### ⏱ {int(row.get('elapsed_minutes') or 0)}分後　`{row.get('category') or 'その他'}`")
    if row.get("content"):
        st.write(row["content"])
    if row.get("question"):
        st.info(f"**出題した問題**\n\n{row['question']}")
    if row.get("updated_at"):
        st.caption(f"最終更新: {format_date(row['updated_at'])}")


def lesson_timeline(client: Client) -> None:
    st.subheader("授業タイムライン")
    try:
        lessons = fetch_rows(client, "lessons")
    except Exception as exc:
        show_firestore_error("授業データを取得", exc)
        return
    lessons.sort(key=lambda row: (row.get("lesson_date") or "", str(row.get("created_at") or "")), reverse=True)
    if not lessons:
        st.info("授業がまだありません。「授業設定・新規」タブから作成してください。")
        return
    unit_tags = sorted({lesson.get("unit") for lesson in lessons if lesson.get("unit")})
    selected_units = st.multiselect("単元タグで絞り込み", unit_tags, placeholder="単元名を検索・選択", key="timeline_units")
    filtered_lessons = [lesson for lesson in lessons if not selected_units or lesson.get("unit") in selected_units]
    if not filtered_lessons:
        st.info("選択した単元タグに該当する授業はありません。")
        return
    labels = {f"{lesson_label(lesson)}（{lesson['id'][-6:]}）": lesson for lesson in filtered_lessons}
    lesson = labels[st.selectbox("表示する授業", list(labels), key="timeline_lesson")]
    if lesson.get("memo"):
        st.caption(f"授業メモ: {lesson['memo']}")
    try:
        cards = fetch_lesson_cards(client, lesson["id"])
    except Exception as exc:
        show_firestore_error("タイムラインを取得", exc)
        return
    export_data = {"lesson": {key: value for key, value in lesson.items() if key != "id"}, "timeline_cards": [{key: value for key, value in card.items() if key != "id"} for card in cards]}
    st.download_button("この授業をJSONで書き出す", data=json.dumps(export_data, ensure_ascii=False, default=str, indent=2), file_name=f"lesson_timeline_{lesson.get('lesson_date', 'record')}.json", mime="application/json")
    st.caption(f"{len(cards)} 件のカード")
    if not cards:
        st.info("まだカードはありません。下の「カードを追加」から記録を始めましょう。")
    for card in cards:
        elapsed = int(card.get("elapsed_minutes") or 0)
        category = card.get("category") or "その他"
        preview = (card.get("content") or "").replace("\n", " ")[:40]
        with st.expander(f"⏱ {elapsed}分後｜{category}｜{preview}"):
            timeline_card(card)
            with st.popover("削除", icon="🗑️"):
                st.warning("このタイムラインカードを削除します。元に戻せません。")
                if st.button("このカードを削除する", type="primary", key=f"delete_lesson_card_{card['id']}"):
                    delete_document(client, "lesson_cards", card["id"], "タイムラインカード")
    st.divider()
    st.subheader("カードを追加")
    with st.form("lesson_card_form", clear_on_submit=True):
        elapsed_minutes = st.number_input("経過時間（分） *", min_value=0, step=1, help="授業開始から何分後かを入力します。")
        category = st.selectbox("カテゴリ", DEFAULT_LESSON_CATEGORIES)
        content = st.text_area("授業内容 *", placeholder="例：現在完了の用法を解説し、例文を確認した。")
        question = st.text_area("出題した問題", placeholder="例：次の文を現在完了形に書き換えなさい。")
        submitted = st.form_submit_button("カードをタイムラインに追加", type="primary")
    if submitted:
        if not content.strip():
            st.warning("「授業内容」は必須です。")
        else:
            now = datetime.now(timezone.utc)
            try:
                client.collection("lesson_cards").add({"lesson_id": lesson["id"], "elapsed_minutes": int(elapsed_minutes), "category": category, "content": content.strip(), "question": question.strip(), "created_at": now, "updated_at": now})
                st.success("タイムラインカードを保存しました。")
                st.rerun()
            except Exception as exc:
                show_firestore_error("タイムラインカードを保存", exc)
    if cards:
        st.divider()
        st.subheader("カードを編集")
        card_labels = {f"{int(card.get('elapsed_minutes') or 0)}分後｜{card.get('category') or 'その他'}｜{card.get('content', '')[:24]}（{card['id'][-6:]}）": card for card in cards}
        selected_card = card_labels[st.selectbox("編集するカード", list(card_labels), key="edit_lesson_card")]
        with st.form(f"edit_lesson_card_form_{selected_card['id']}"):
            elapsed_minutes = st.number_input("経過時間（分）", min_value=0, step=1, value=int(selected_card.get("elapsed_minutes") or 0))
            category = st.selectbox("カテゴリ", DEFAULT_LESSON_CATEGORIES, index=DEFAULT_LESSON_CATEGORIES.index(selected_card["category"]) if selected_card.get("category") in DEFAULT_LESSON_CATEGORIES else len(DEFAULT_LESSON_CATEGORIES) - 1)
            content = st.text_area("授業内容", value=selected_card.get("content", ""))
            question = st.text_area("出題した問題", value=selected_card.get("question", ""))
            submitted = st.form_submit_button("変更を保存する")
        if submitted:
            if not content.strip():
                st.warning("「授業内容」は必須です。")
            else:
                try:
                    client.collection("lesson_cards").document(selected_card["id"]).update({"elapsed_minutes": int(elapsed_minutes), "category": category, "content": content.strip(), "question": question.strip(), "updated_at": datetime.now(timezone.utc)})
                    st.success("タイムラインカードを更新しました。")
                    st.rerun()
                except Exception as exc:
                    show_firestore_error("タイムラインカードを更新", exc)


def build_reading_test_prompt(passage: str, level: str, question_plan: dict[str, int]) -> str:
    return f"""あなたは英語教育のプロフェッショナルな塾講師・教材作成者です。
以下の英語長文を元に、指定された条件に従って生徒用の小テストと解答・解説を作成してください。

【対象レベル】
{level}

【作成する問題種別と問題数】
{chr(10).join(f"- {question_type}: {count}問" for question_type, count in question_plan.items())}

【問題数】
計 {sum(question_plan.values())} 問

【出力形式の指定】
必ず以下の2つのセクションに明確に分けて出力してください。

### 【問題用紙】
- 生徒にそのまま印刷・配布できるフォーマットにし、最初に【本文】を掲載してください。
- 「語順並べ替え問題」が含まれる場合、本文中の該当英文を空欄（①, ②など）に置き換え、空欄の直後にランダムな単語プールを角括弧で埋め込んでください。（例: ① [ study / hard / to / dynamic / should ]）。本文の外に同じ問題を再掲しないでください。
- 「和文英訳問題」が含まれる場合、本文中の該当英文を空欄（①, ②など）に置き換え、空欄の直後にその英文の日本語訳を（　）内に埋め込んでください。本文の外に同じ問題を再掲しないでください。
- 各問題の指示文は日本語で分かりやすく記述してください。

### 【解答・解説】
- 各問題の正解を明記してください。
- 語順並べ替え問題については、完成文とその和訳、なぜその語順になるのかの文法ポイント（構文、文型、熟語など）を中高生に分かりやすく解説してください。
- その他の問題についても、根拠となる本文の該当箇所や文法・語彙のポイントを解説に含めてください。

【英語長文】
{passage}
"""


def split_reading_test_sections(generated_text: str) -> tuple[str, str]:
    question_marker, answer_marker = "### 【問題用紙】", "### 【解答・解説】"
    if question_marker in generated_text and answer_marker in generated_text:
        _, after_question = generated_text.split(question_marker, 1)
        question_paper, answer_key = after_question.split(answer_marker, 1)
        return question_paper.strip(), answer_key.strip()
    return generated_text.strip(), "Geminiの出力を2つのセクションへ分割できませんでした。全文を問題用紙として表示しています。"


def make_word_document(title: str, content: str) -> bytes:
    """Create a downloadable Word document without writing generated text to disk."""
    from docx import Document

    document = Document()
    document.add_heading(title, level=0)
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        else:
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def reading_test_generator() -> None:
    st.subheader("長文読解テスト・ジェネレータ")
    st.caption("Geminiを利用して、英語長文から問題用紙と解答・解説を作成します。")
    passage = st.text_area("長文テキスト入力", height=260, placeholder="英語長文を貼り付けてください。", key="reading_passage")
    level = st.selectbox("対象レベル", READING_LEVELS, index=2)
    selected_question_types = st.multiselect("問題種別を選択", READING_QUESTION_TYPES, default=["語順並べ替え問題", "内容一致問題", "空所補充問題"])
    st.markdown("#### 選択した問題種別ごとの問題数")
    st.caption("選択した種別だけ問題数を設定します。合計は1〜10問にしてください。")
    default_counts = {"語順並べ替え問題": 2, "内容一致問題": 2, "空所補充問題": 1}
    question_counts = {question_type: st.number_input(f"{question_type}（問）", min_value=0, max_value=10, value=default_counts.get(question_type, 0), step=1, key=f"reading_count_{question_type}") for question_type in selected_question_types}
    question_plan = {question_type: int(count) for question_type, count in question_counts.items() if count > 0}
    total_questions = sum(question_plan.values())
    st.info(f"合計 {total_questions} 問")
    if st.button("テストを自動生成する", type="primary", use_container_width=True):
        if not passage.strip():
            st.warning("英語長文を入力してください。")
        elif not question_plan:
            st.warning("問題種別ごとの問題数を1問以上にしてください。")
        elif total_questions > 10:
            st.warning("問題数の合計は10問以下にしてください。")
        else:
            api_key = secret_or_env("GEMINI_API_KEY")
            if not api_key:
                st.error("GEMINI_API_KEY が未設定です。Streamlit Secrets または環境変数に設定してください。")
            else:
                try:
                    from google import genai

                    model = secret_or_env("GEMINI_MODEL") or "gemini-2.5-flash"
                    with st.spinner("Geminiがテストを作成中..."):
                        # Keep the SDK client alive until the response has been received.
                        gemini_client = genai.Client(api_key=api_key)
                        response = gemini_client.models.generate_content(model=model, contents=build_reading_test_prompt(passage.strip(), level, question_plan))
                    generated_text = (response.text or "").strip()
                    if not generated_text:
                        raise ValueError("Geminiからテキストが返されませんでした。")
                    question_paper, answer_key = split_reading_test_sections(generated_text)
                    st.session_state["reading_test_result"] = {"question_paper": question_paper, "answer_key": answer_key, "full_text": generated_text}
                except ImportError:
                    st.error("google-genai が未インストールです。requirements.txt を更新して再デプロイしてください。")
                except Exception as exc:
                    st.error(f"Geminiによるテスト生成に失敗しました: {exc}")
    result = st.session_state.get("reading_test_result")
    if not result:
        return
    question_tab, answer_tab = st.tabs(["問題用紙", "解答・解説"])
    with question_tab:
        st.code(result["question_paper"], language=None)
        st.download_button("問題用紙（Word）をダウンロード", make_word_document("長文読解テスト：問題用紙", result["question_paper"]), "reading_test_questions.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with answer_tab:
        st.code(result["answer_key"], language=None)
        st.download_button("解答・解説（Word）をダウンロード", make_word_document("長文読解テスト：解答・解説", result["answer_key"]), "reading_test_answers.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def main() -> None:
    if "current_page" not in st.session_state:
        st.session_state.current_page = "メインメニュー"

    page = st.session_state.current_page
    if page == "メインメニュー":
        st.title("📚 My English Dictionary")
        st.caption("学習の記録・教材づくりを、ひとつの場所で。")
        st.divider()
        st.subheader("メインメニュー")
        st.caption("使いたい機能のアイコンをクリックしてください。")
        menu_items = [
            ("💬", "引用", "映画や小説の引用を記録・検索"),
            ("🌱", "語源", "単語と語源をつなげて記録"),
            ("🕒", "授業タイムライン", "授業の流れと出題内容を可視化"),
            ("📝", "長文読解テスト", "英語長文からテストを自動生成"),
        ]
        for start in range(0, len(menu_items), 2):
            columns = st.columns(2)
            for column, (icon, label, description) in zip(columns, menu_items[start:start + 2]):
                with column:
                    st.markdown(f"### {icon} {label}")
                    st.caption(description)
                    if st.button(f"{icon} 開く", key=f"open_{label}", use_container_width=True):
                        st.session_state.current_page = label
                        st.rerun()
        return

    home_column, heading_column = st.columns([1, 12])
    with home_column:
        if st.button("⌂", key=f"home_{page}", help="メインメニューへ戻る"):
            st.session_state.current_page = "メインメニュー"
            st.rerun()
    with heading_column:
        st.title(page)

    if page == "長文読解テスト":
        reading_test_generator()
        return

    client = get_client()
    if client is None:
        st.stop()
    if page == "授業タイムライン":
        timeline_tab, settings_tab = st.tabs(["タイムライン", "授業設定・新規"])
        with timeline_tab:
            lesson_timeline(client)
        with settings_tab:
            lesson_settings(client)
        return

    list_tab, form_tab = st.tabs(["一覧・検索", "新規登録"])
    if page == "引用":
        with list_tab:
            quote_list(client)
        with form_tab:
            quote_form(client)
    else:
        with list_tab:
            etymology_list(client)
        with form_tab:
            etymology_form(client)


if __name__ == "__main__":
    main()
